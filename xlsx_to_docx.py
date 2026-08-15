#!/usr/bin/env python3
"""Convert the internship xlsx to a well-formatted docx table."""
import openpyxl
from docx import Document
from docx.shared import Pt, Cm, Inches, RGBColor
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

src = "/opt/xiaonai/tmp_实习岗位.xlsx"
out = "/opt/xiaonai/2026年武汉市大学生实习实训岗位信息登记表.docx"

wb = openpyxl.load_workbook(src)
ws = wb.active

doc = Document()

# Page margins
for section in doc.sections:
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(1.5)
    section.right_margin = Cm(1.5)

# Title
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run("2026年武汉市大学生实习实训岗位信息登记表")
run.bold = True
run.font.size = Pt(16)
run.font.name = "微软雅黑"

# Subtitle line
sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.LEFT
run = sub.add_run("报送学院：（盖章）                           报送人：                    ")
run.font.size = Pt(11)

# Read all data rows from xlsx
rows_data = []
for row in ws.iter_rows(min_row=4, values_only=False):
    vals = [c.value for c in row]
    if vals[0] is not None or vals[1] is not None:
        rows_data.append(vals)

# Filter out completely empty rows and NOTE row
header_data = []
for row in ws.iter_rows(min_row=1, max_row=3, values_only=True):
    header_data.append(list(row))

# Build table: header columns 0-7 (岗位信息) + columns 8-15 (学生填报)
table = doc.add_table(rows=1, cols=16)
table.alignment = WD_TABLE_ALIGNMENT.CENTER
table.style = 'Table Grid'

# Set column widths
col_widths = [Cm(3.5), Cm(2.5), Cm(1.2), Cm(3.0), Cm(1.0), Cm(4.0), Cm(1.2), Cm(2.5),
              Cm(0.8), Cm(1.5), Cm(1.8), Cm(2.0), Cm(1.5), Cm(3.0), Cm(2.0), Cm(1.5)]

# Column headers
headers = ["单位名称", "岗位名称", "学历要求", "岗位工作描述", "招录人数", "岗位要求", "所属地区", "详细地点",
           "序号", "姓名", "学院", "专业", "年级班级", "身份证号", "联系方式", "是否调剂"]

hdr_cells = table.rows[0].cells
for i, h in enumerate(headers):
    hdr_cells[i].text = ""
    p = hdr_cells[i].paragraphs[0]
    run = p.add_run(h)
    run.bold = True
    run.font.size = Pt(8)
    run.font.name = "微软雅黑"
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    # Shading
    from docx.oxml.ns import qn
    shading = hdr_cells[i]._tc.get_or_add_tcPr()
    shading_elm = shading.makeelement(qn('w:shd'), {
        qn('w:fill'): 'D9E2F3',
        qn('w:val'): 'clear'
    })
    shading.append(shading_elm)

# Data rows
unit_colspan = {}  # track which rows need colspan for unit name
current_unit = None
unit_start = None

for idx, row_vals in enumerate(rows_data):
    row_cells = table.add_row().cells
    
    for col_idx in range(16):
        val = row_vals[col_idx] if col_idx < len(row_vals) else ""
        if val is None:
            val = ""
        val = str(val).strip()
        
        row_cells[col_idx].text = ""
        p = row_cells[col_idx].paragraphs[0]
        run = p.add_run(val)
        run.font.size = Pt(7.5)
        run.font.name = "微软雅黑"
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

# Merge empty unit name cells upward
# (skip complex merge for now, simple representation)

# Add note at bottom
note = doc.add_paragraph()
note.add_run("\n备注：").bold = True
note.add_run("学生信息填写在申报的岗位对应栏，入选学生需确保2026年7月7日-8月7日期间可按时到岗参与工作。请各学院于6月5日前提交。")

doc.save(out)
print(f"OK: {out} ({os.path.getsize(out)} bytes)")
