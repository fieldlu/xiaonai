#!/usr/bin/env python3
"""Create a simplified-vocabulary version of Detroit-逐词翻译.docx for junior-high level."""

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import copy

# ===== SIMPLIFIED ENGLISH TEXT (初中生水平) =====
# Original -> Simplified mapping for each sentence
simplified_sentences = [
    (
        "1. Detroit, the largest city in the midwestern state of Michigan, holds a unique and indispensable position in American industrial history.",
        "1. Detroit, the biggest city in the midwestern state of Michigan, has a special and very important place in American industrial history."
    ),
    (
        '2. Globally recognized as the "Motor City," Detroit is intrinsically synonymous with the birth and evolution of the automobile industry.',
        '2. Known around the world as the "Motor City," Detroit is deeply connected to the start and growth of the car industry.'
    ),
    (
        "3. It was here that early pioneers perfected the assembly line, fundamentally transforming modern manufacturing processes and shaping global transportation.",
        "3. It was here that early pioneers made the assembly line perfect, greatly changing how things are made today and shaping transportation around the world."
    ),
    (
        "4. Although the city faced severe economic downturns in the late 20th century, it is currently undergoing a remarkable urban revitalization.",
        "4. Although the city had very hard economic times in the late 1900s, it is now going through an amazing city rebirth."
    ),
    (
        "5. Detroit is aggressively diversifying its economic portfolio, investing heavily in green technology, smart mobility, and urban renewal projects.",
        "5. Detroit is actively trying to grow many different parts of its economy, putting a lot of money into green technology, smart transportation, and city renewal projects."
    ),
    (
        "6. Beyond its industrial legacy, the city also possesses a profound cultural heritage, most notably as the birthplace of Motown Records, which profoundly revolutionized American music.",
        "6. Besides its history as an industrial city, Detroit also has a deep cultural background, most famously as the birthplace of Motown Records, which deeply changed American music."
    ),
    (
        "7. Ultimately, Detroit stands as a powerful testament to resilience and industrial ingenuity.",
        "7. In the end, Detroit stands as a strong example of the ability to bounce back and creative thinking in industry."
    ),
    (
        "8. Its ongoing transformation from a historical manufacturing hub to a modern, diversified metropolis demonstrates its enduring capacity for reinvention.",
        "8. Its ongoing change from an old manufacturing center to a modern, diverse big city shows its lasting ability to reinvent itself."
    ),
]

# Word-level simplified versions (for word-by-word translation section)
word_translations = [
    # Sentence 1
    ("holds a unique and indispensable position",
     "holds a special and very important place",
     "拥有一个独特且不可或缺的地位 -> 拥有一个特殊且非常重要的位置"),
    # Sentence 2
    ("Globally recognized as",
     "Known around the world as",
     "全球公认 -> 全世界都知道"),
    ("intrinsically synonymous with",
     "deeply connected to",
     "本质上同义的 -> 紧密相连的"),
    ("birth and evolution",
     "start and growth",
     "诞生和演变 -> 开始和发展"),
    # Sentence 3
    ("early pioneers",
     "early pioneers",
     "早期的先驱们 -> 早期的开拓者"),
    ("fundamentally transforming",
     "greatly changing",
     "根本地改变 -> 极大地改变"),
    # Sentence 4
    ("severe economic downturns",
     "very hard economic times",
     "严重的经济衰退 -> 非常艰难的经济时期"),
    ("remarkable urban revitalization",
     "amazing city rebirth",
     "引人注目的城市复兴 -> 令人惊叹的城市重生"),
    # Sentence 5
    ("aggressively diversifying",
     "actively growing different parts of",
     "积极地多样化 -> 积极地拓展不同领域"),
    ("economic portfolio",
     "economy / different parts of the economy",
     "经济结构 -> 经济的各个方面"),
    ("smart mobility",
     "smart transportation",
     "智能出行 -> 智能交通"),
    # Sentence 6
    ("profound cultural heritage",
     "deep cultural background",
     "深厚的文化遗产 -> 深厚的文化底蕴"),
    ("profoundly revolutionized",
     "deeply changed",
     "深刻地彻底改变 -> 深刻地改变"),
    # Sentence 7
    ("a powerful testament to",
     "a strong example of",
     "一个强大的证明 -> 一个有力的例证"),
    ("resilience",
     "the ability to bounce back",
     "韧性 -> 恢复力/反弹能力"),
    ("industrial ingenuity",
     "creative thinking in industry",
     "工业独创性 -> 工业创造力"),
    # Sentence 8
    ("ongoing transformation",
     "ongoing change",
     "正在进行的转型 -> 正在进行的变化"),
    ("historical manufacturing hub",
     "old manufacturing center",
     "历史悠久的制造业中心 -> 老制造中心"),
    ("diversified metropolis",
     "diverse big city",
     "多元化的大都市 -> 多元化的大城市"),
    ("enduring capacity for reinvention",
     "lasting ability to reinvent itself",
     "持久的重新创造能力 -> 持久的自我革新能力"),
]

# ===== BUILD DOCUMENT =====
doc = Document()

# Title
title = doc.add_heading('Essay 2: Detroit — 初中生版逐词翻译', level=1)
title.runs[0].font.color.rgb = RGBColor(0x1a, 0x5c, 0x8a)

note = doc.add_paragraph()
note_run = note.add_run('说明：把原文难词换成简单的同义词，初中生水平也能看懂。保留逐词对照结构，方便记忆。')
note_run.font.size = Pt(11)
note_run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
note_run.italic = True

doc.add_paragraph()

# SECTION 1: 复杂词汇替换对照表
h1 = doc.add_heading('一、复杂词汇替换对照表', level=2)

table = doc.add_table(rows=1, cols=3)
table.style = 'Light Grid Accent 1'
hdr_cells = table.rows[0].cells
hdr_cells[0].text = '原词/词组'
hdr_cells[1].text = '简单替换'
hdr_cells[2].text = '中文说明'

for orig, simp, zh in word_translations:
    row_cells = table.add_row().cells
    row_cells[0].text = orig
    row_cells[1].text = simp
    row_cells[2].text = zh

doc.add_paragraph()

# SECTION 2: 完整原文 + 简化版 + 逐词翻译对照
h2 = doc.add_heading('二、完整原文 + 简化版 + 逐词翻译对照', level=2)

for orig, simp in simplified_sentences:
    # Original
    p_orig = doc.add_paragraph()
    p_orig.add_run(orig).bold = True
    
    # Simplified
    p_simp = doc.add_paragraph()
    run_label = p_simp.add_run('【简化】')
    run_label.font.color.rgb = RGBColor(0x00, 0x8a, 0x00)
    run_label.font.size = Pt(10)
    run_label.bold = True
    p_simp.add_run(simp)
    
    doc.add_paragraph()  # spacing

# Save
out_path = '/opt/xiaonai/exports/Detroit-初中生版逐词翻译.docx'
doc.save(out_path)
print(f'Saved to {out_path}')
