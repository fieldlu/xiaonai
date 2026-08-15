#!/usr/bin/env python3
"""XiaoNai Literature Export Tool — Search papers + export as formatted .docx.

Modes:
  python3 tools/docx_export_helper.py search <query> [--rows N] -o <output.docx>
  python3 tools/docx_export_helper.py from-json <papers.json> -o <output.docx> [--query "..."]

Mode 'search': internally calls scholar_search, then generates .docx (recommended).
Mode 'from-json': reads existing papers JSON, generates .docx (debug/regeneration).

Exit codes:
  0 = success
  1 = no results / empty
  2 = argument error
  3 = generation failed
"""
import sys, os, json, subprocess, re, argparse, tempfile
from datetime import datetime
from pathlib import Path

# python-docx imports (guarded for help/validation runs)
try:
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor, Cm, Emu
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml.ns import qn, nsdecls
    from docx.oxml import parse_xml
    from docx.opc.constants import RELATIONSHIP_TYPE as RT
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

# Ensure output dir
_SCRIPT_DIR = Path(__file__).resolve().parent
OUTPUT_DIR = _SCRIPT_DIR / "exports"
OUTPUT_DIR.mkdir(parents=True, exist_ok=True)

# OpenAlex data source
SOURCES = ["OpenAlex (2.44亿+学术作品)"]


# ════════════════════════════════════════════════════════════════
#  Paper search — delegate to scholar_search.py
# ════════════════════════════════════════════════════════════════

def search_papers(query: str, rows: int = 10) -> dict:
    """Run scholar_search.py and return parsed JSON result."""
    scholar = (_SCRIPT_DIR / ".." / "search" / "scholar_search.py").resolve()
    if not scholar.exists():
        # fallback to PATH
        scholar = "scholar_search.py"

    result = subprocess.run(
        [sys.executable, str(scholar), "search", query, "--rows", str(rows)],
        capture_output=True, text=True, timeout=60
    )
    if result.returncode != 0:
        err = result.stderr.strip() or "unknown error"
        raise RuntimeError(f"scholar_search failed: {err}")

    data = json.loads(result.stdout)
    return data


def load_papers_json(path: str) -> dict:
    """Load papers from a JSON file."""
    with open(path, "r", encoding="utf-8") as f:
        data = json.load(f)
    return data


# ════════════════════════════════════════════════════════════════
#  DOCX generation
# ════════════════════════════════════════════════════════════════

def build_docx(papers: list, query: str, total: int, sources: list):
    """Generate a beautifully formatted literature survey .docx."""
    if not DOCX_AVAILABLE:
        raise RuntimeError("python-docx not installed")

    doc = Document()

    # ── Page setup: A4, 1-inch margins ──
    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(2.54)
    section.right_margin = Cm(2.54)

    # ── Default style ──
    style = doc.styles["Normal"]
    style.font.name = "Arial"
    style.font.size = Pt(11)
    style.paragraph_format.space_after = Pt(6)
    style.paragraph_format.line_spacing = 1.15

    # ════════════════════════════════════
    #  COVER PAGE
    # ════════════════════════════════════

    # Spacer
    for _ in range(4):
        doc.add_paragraph()

    # Main title
    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_para.add_run(f"{query}\n文献调研报告")
    title_run.font.size = Pt(28)
    title_run.font.bold = True
    title_run.font.color.rgb = RGBColor(0x1A, 0x3C, 0x6E)

    doc.add_paragraph()  # spacer

    # Subtitle — metadata
    meta_lines = [
        f"检索关键词: {query}",
        f"检索日期: {datetime.now().strftime('%Y年%m月%d日')}",
        f"检索结果: 共 {total} 篇论文，覆盖 {len(sources)} 个学术数据库",
        f"生成工具: 小奈 QQ 机器人 · OpenAlex API",
    ]
    for line in meta_lines:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(line)
        r.font.size = Pt(12)
        r.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

    # Decorative line
    doc.add_paragraph()
    deco = doc.add_paragraph()
    deco.alignment = WD_ALIGN_PARAGRAPH.CENTER
    dr = deco.add_run("─" * 50)
    dr.font.color.rgb = RGBColor(0x2E, 0x75, 0xB6)
    dr.font.size = Pt(10)

    # Page break
    doc.add_page_break()

    # ════════════════════════════════════
    #  SEARCH STRATEGY
    # ════════════════════════════════════

    h1 = doc.add_heading("一、检索策略", level=1)
    _set_heading_font(h1, size=Pt(16), color=RGBColor(0x1A, 0x3C, 0x6E))

    info_table = doc.add_table(rows=3, cols=2)
    info_table.style = "Table Grid"
    info_table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Set column widths
    for cell in info_table.columns[0].cells:
        cell.width = Cm(4)
    for cell in info_table.columns[1].cells:
        cell.width = Cm(12)

    info_data = [
        ("检索关键词", query),
        ("检索日期", datetime.now().strftime('%Y-%m-%d %H:%M')),
        ("数据来源", "、".join(_format_source(s) for s in sources)),
    ]
    for i, (k, v) in enumerate(info_data):
        _set_cell(info_table.rows[i].cells[0], k, bold=True, bg="D5E8F0")
        _set_cell(info_table.rows[i].cells[1], v)

    doc.add_paragraph()  # spacer

    # ════════════════════════════════════
    #  PAPER OVERVIEW TABLE
    # ════════════════════════════════════

    h2 = doc.add_heading("二、文献清单", level=1)
    _set_heading_font(h2, size=Pt(16), color=RGBColor(0x1A, 0x3C, 0x6E))

    p_count = doc.add_paragraph(f"共检索到 {total} 篇相关论文，以下展示前 {len(papers)} 篇：")
    p_count.runs[0].font.size = Pt(10)
    p_count.runs[0].font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    # Create table
    headers = ["#", "标题", "第一作者", "年份", "来源", "引用"]
    col_widths = [Cm(1), Cm(8.5), Cm(3), Cm(1.5), Cm(2.5), Cm(1.5)]

    table = doc.add_table(rows=1 + len(papers), cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Set column widths
    for i, w in enumerate(col_widths):
        for row in table.rows:
            row.cells[i].width = w

    # Header row
    for i, h in enumerate(headers):
        _set_cell(table.rows[0].cells[i], h, bold=True, bg="1A3C6E", font_color=RGBColor(0xFF, 0xFF, 0xFF), size=Pt(10))

    # Data rows
    for idx, paper in enumerate(papers, 1):
        row = table.rows[idx]
        authors = _parse_authors(paper.get("authors", []))
        first_author = authors[0] + "等" if len(authors) > 1 else (authors[0] if authors else "—")
        year = str(paper.get("year", "")) or "—"
        source = _format_source(paper.get("source", ""))
        citations = str(paper.get("citations", 0) or 0)
        title = paper.get("title", "?")
        doi = paper.get("doi", "")

        # Row data
        _set_cell(row.cells[0], str(idx), align=WD_ALIGN_PARAGRAPH.CENTER, size=Pt(9))
        # Title only (DOI has its own column)
        _set_title_cell(row.cells[1], title, doi, size=Pt(9))
        _set_cell(row.cells[2], first_author, size=Pt(9))
        _set_cell(row.cells[3], year, align=WD_ALIGN_PARAGRAPH.CENTER, size=Pt(9))
        _set_cell(row.cells[4], source, size=Pt(9))
        _set_cell(row.cells[5], citations, align=WD_ALIGN_PARAGRAPH.CENTER, size=Pt(9))

        # Alternating row color
        if idx % 2 == 0:
            for cell in row.cells:
                _shade_cell(cell, "F2F7FB")

    doc.add_page_break()

    # ════════════════════════════════════
    #  PAPER DETAILS
    # ════════════════════════════════════

    h3 = doc.add_heading("三、论文详情", level=1)
    _set_heading_font(h3, size=Pt(16), color=RGBColor(0x1A, 0x3C, 0x6E))

    for idx, paper in enumerate(papers, 1):
        # Paper number header
        h4 = doc.add_heading(f"第 {idx} 篇", level=2)
        _set_heading_font(h4, size=Pt(13), color=RGBColor(0x2E, 0x75, 0xB6))

        # Title with DOI link
        title = paper.get("title", "?")
        doi = paper.get("doi", "")
        title_p = doc.add_paragraph()
        title_p.paragraph_format.space_after = Pt(4)
        tr = title_p.add_run(f"标题: ")
        tr.bold = True
        tr.font.size = Pt(11)

        if doi:
            from docx.opc.constants import RELATIONSHIP_TYPE as RT
            # Add hyperlink relationship
            doi_url = f"https://doi.org/{doi}"
            link_id = doc.part.relate_to(doi_url, RT.HYPERLINK, is_external=True)
            hyperlink = parse_xml(
                f'<w:hyperlink {nsdecls("w")} r:id="{link_id}" {nsdecls("r")}>'
                f'  <w:r><w:rPr><w:rStyle w:val="Hyperlink"/><w:sz w:val="22"/></w:rPr>'
                f'  <w:t xml:space="preserve">{_escape_xml(title)}</w:t></w:r>'
                f'</w:hyperlink>'
            )
            title_p._p.append(hyperlink)
        else:
            tr2 = title_p.add_run(title)
            tr2.font.size = Pt(11)

        # Metadata table for each paper
        detail_table = doc.add_table(rows=5, cols=2)
        detail_table.style = "Table Grid"
        for cell in detail_table.columns[0].cells:
            cell.width = Cm(3)
        for cell in detail_table.columns[1].cells:
            cell.width = Cm(13)

        authors_str = "、".join(_parse_authors(paper.get("authors", []))) or "—"
        year = str(paper.get("year", "")) or "—"
        source = _format_source(paper.get("source", ""))
        citations_str = str(paper.get("citations", 0) or 0) + " 次引用"
        journal = paper.get("journal", "") or "—"

        detail_data = [
            ("作者", authors_str),
            ("年份", year),
            ("期刊", journal),
            ("来源数据库", source),
            ("引用", citations_str),
        ]
        for i, (k, v) in enumerate(detail_data):
            _set_cell(detail_table.rows[i].cells[0], k, bold=True, bg="E8F0FE", size=Pt(10))
            _set_cell(detail_table.rows[i].cells[1], v, size=Pt(10))

        # DOI row (if available)
        if doi:
            doi_row = detail_table.add_row()
            _set_cell(doi_row.cells[0], "DOI", bold=True, bg="E8F0FE", size=Pt(10))
            doi_url = f"https://doi.org/{doi}"
            doi_p = doi_row.cells[1].paragraphs[0]
            doi_run = doi_p.add_run(doi_url)
            doi_run.font.size = Pt(9)
            doi_run.font.color.rgb = RGBColor(0x2E, 0x75, 0xB6)

        # Abstract section
        abstract = paper.get("abstract", "") or ""
        if abstract:
            abs_title = doc.add_paragraph()
            abs_title.paragraph_format.space_before = Pt(6)
            abs_title.paragraph_format.space_after = Pt(2)
            atr = abs_title.add_run("摘要: ")
            atr.bold = True
            atr.font.size = Pt(10)
            atr.font.color.rgb = RGBColor(0x1A, 0x3C, 0x6E)

            abs_text = doc.add_paragraph()
            abs_text.paragraph_format.space_after = Pt(4)
            abs_text.paragraph_format.first_line_indent = Cm(0.5)
            ar = abs_text.add_run(abstract[:600])
            ar.font.size = Pt(9.5)
            ar.font.color.rgb = RGBColor(0x33, 0x33, 0x33)

        # Separator
        sep = doc.add_paragraph()
        sep.paragraph_format.space_before = Pt(2)
        sep.paragraph_format.space_after = Pt(8)
        sr = sep.add_run("─" * 60)
        sr.font.size = Pt(8)
        sr.font.color.rgb = RGBColor(0xCC, 0xCC, 0xCC)

    # ════════════════════════════════════
    #  SOURCE DECLARATION
    # ════════════════════════════════════

    doc.add_page_break()
    h5 = doc.add_heading("四、数据来源说明", level=1)
    _set_heading_font(h5, size=Pt(16), color=RGBColor(0x1A, 0x3C, 0x6E))

    source_intro = (
        "本报告数据来源于 OpenAlex API (https://openalex.org)，"
        "覆盖全球 2.44 亿+学术作品的开放元数据，包括论文、作者、期刊、引用关系和摘要索引。"
    )
    doc.add_paragraph(source_intro)

    # Source list
    for s in SOURCES:
        bp = doc.add_paragraph(style="List Bullet")
        bp.add_run(s).font.size = Pt(10)

    doc.add_paragraph()
    disclaimer = doc.add_paragraph(
        "声明: 本报告仅基于上述数据库的公开检索结果自动生成，不保证文献的完整性和准确性。"
        "数据版权归各数据库所有。生成的文档仅供学术参考，请通过正式渠道获取全文。"
    )
    disclaimer.runs[0].font.size = Pt(9)
    disclaimer.runs[0].font.color.rgb = RGBColor(0x88, 0x88, 0x88)
    disclaimer.runs[0].italic = True

    return doc


# ════════════════════════════════════
#  Helper functions
# ════════════════════════════════════

def _parse_authors(authors_field):
    """Parse authors field (string or list) into a clean list of names."""
    if not authors_field:
        return []
    if isinstance(authors_field, list):
        raw = authors_field
    elif isinstance(authors_field, str):
        raw = [a.strip() for a in authors_field.replace("，", ",").split(",") if a.strip()]
    else:
        return []
    # Clean up: remove institution names (Chinese universities etc.)
    cleaned = []
    for name in raw:
        name = name.strip()
        # Skip institution names (contain 大学/学院/研究所/中心 etc.)
        if any(kw in name for kw in ["大学", "学院", "研究所", "中心", "实验室", "公司"]):
            continue
        # Skip empty or single-char
        if len(name) < 2:
            continue
        # Fix Chinese name order:
        # "镇宗 王" (given+surname Western order) -> "王镇宗"
        # "唐 杰" (space artifact) -> "唐杰"
        parts = name.split()
        if len(parts) == 2:
            is_cjk = all('一' <= c <= '鿿' for c in parts[0] + parts[1])
            if is_cjk:
                name = parts[0] + parts[1]  # remove space, keep order
        cleaned.append(name)
    return cleaned


def _format_source(source):
    """Capitalize source name for display."""
    name_map = {
        "crossref": "CrossRef",
        "openalex": "OpenAlex",
        "semantic_scholar": "Semantic Scholar",
        "pubmed": "PubMed",
        "arxiv": "arXiv",
        "core": "CORE",
        "europe_pmc": "Europe PMC",
        "doaj": "DOAJ",
        "dblp": "dblp",
    }
    return name_map.get(source.lower(), source) if source else "—"


def _build_doi_url(doi):
    """Build DOI URL and return (url_text, url)."""
    if not doi:
        return None, None
    return f"https://doi.org/{doi}", doi


def _set_heading_font(heading, size=None, color=None):
    """Set heading formatting. Size defaults to 16pt, color to dark blue."""
    if size is None:
        size = Pt(16)
    if color is None:
        color = RGBColor(0x1A, 0x3C, 0x6E)
    for run in heading.runs:
        run.font.size = size
        run.font.color.rgb = color


def _set_cell(cell, text, bold=False, bg=None, font_color=None, size=None,
              align=None):
    """Set cell text and formatting."""
    if size is None:
        size = Pt(10)
    if align is None:
        align = WD_ALIGN_PARAGRAPH.LEFT
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = align
    run = p.add_run(str(text))
    run.font.size = size
    run.font.name = "Arial"
    if bold:
        run.bold = True
    if font_color:
        run.font.color.rgb = font_color

    # Cell vertical alignment
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    vAlign = parse_xml(f'<w:vAlign {nsdecls("w")} w:val="center"/>')
    tcPr.append(vAlign)

    # Background shading
    if bg:
        _shade_cell(cell, bg)


def _shade_cell(cell, color):
    """Apply background shading to a table cell."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shading = parse_xml(
        f'<w:shd {nsdecls("w")} w:fill="{color}" w:val="clear"/>'
    )
    tcPr.append(shading)


def _set_title_cell(cell, title, doi=None, size=None):
    """Set title cell with title text only. DOI shown in its own column."""
    if size is None:
        size = Pt(9)
    cell.text = ""
    p = cell.paragraphs[0]
    run = p.add_run(title[:120])
    run.font.size = size
    run.font.name = "Arial"


def _escape_xml(text: str) -> str:
    """Escape text for XML."""
    return text.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;").replace('"', "&quot;")


# ════════════════════════════════════
#  Main entry point
# ════════════════════════════════════

def main():
    parser = argparse.ArgumentParser(description="XiaoNai Literature Export Tool")
    sub = parser.add_subparsers(dest="mode", required=True)

    # Mode: search
    p_search = sub.add_parser("search", help="Search papers + export to docx")
    p_search.add_argument("query", help="Search query")
    p_search.add_argument("--rows", type=int, default=10, help="Number of results (default: 10)")
    p_search.add_argument("-o", "--output", required=True, help="Output .docx path")

    # Mode: from-json
    p_json = sub.add_parser("from-json", help="Generate docx from existing JSON")
    p_json.add_argument("json_file", help="Path to papers JSON file")
    p_json.add_argument("-o", "--output", required=True, help="Output .docx path")
    p_json.add_argument("--query", default="文献检索", help="Search query (for title)")

    args = parser.parse_args()

    try:
        if args.mode == "search":
            # Step 1: Search
            data = search_papers(args.query, args.rows)
            papers = data.get("results", data.get("papers", []))
            total = data.get("total", len(papers))
            sources = data.get("sources", SOURCES)
            query = data.get("query", args.query)

            if not papers:
                print(f"No results for: {args.query}", file=sys.stderr)
                sys.exit(1)

            # Step 2: Generate docx
            doc = build_docx(papers, query, total, sources)
            output = args.output

        elif args.mode == "from-json":
            # Load from file
            data = load_papers_json(args.json_file)
            papers = data.get("results", data.get("papers", data if isinstance(data, list) else []))
            total = data.get("total", len(papers))
            sources = data.get("sources", SOURCES)
            query = args.query

            if not papers:
                print(f"No papers in JSON: {args.json_file}", file=sys.stderr)
                sys.exit(1)

            doc = build_docx(papers, query, total, sources)
            output = args.output

        # Ensure output directory exists
        # CRITICAL: resolve relative paths against script dir (not CWD)
        # Agent runs from ~/.openclaw/workspace/, not the project root
        output_path = Path(output)
        if not output_path.is_absolute():
            output_path = _SCRIPT_DIR / output_path
        output_path.parent.mkdir(parents=True, exist_ok=True)

        # Save
        doc.save(str(output_path))
        size = output_path.stat().st_size
        print(f"Created: {output} ({size} bytes)")

    except subprocess.TimeoutExpired:
        print("ERROR: scholar_search timed out (60s)", file=sys.stderr)
        sys.exit(3)
    except json.JSONDecodeError as e:
        print(f"ERROR: Invalid JSON: {e}", file=sys.stderr)
        sys.exit(3)
    except Exception as e:
        print(f"ERROR: {e}", file=sys.stderr)
        sys.exit(3)


if __name__ == "__main__":
    main()
